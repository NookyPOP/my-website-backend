import datetime
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches


def create_excel():
    wb = Workbook()
    ws = wb.active
    ws["A1"] = 11
    ws.append([1, 2, 3, 4, 5, 6])
    ws["A2"] = datetime.datetime.now()
    ws["B1"] = "jack"
    wb.save("python-practice/sample.xlsx")


def create_word():
    document = Document()
    document.add_heading("Document title", 0)
    p = document.add_paragraph("A plain paragraph having some")
    p.add_run("bold").bold = True
    p.add_run(" and some boies")
    p.add_run("and italic.").italic = True

    document.add_heading("Heading Level1", level=1)
    document.add_paragraph("Intense quote", style="Intense Quote")
    document.add_paragraph("first item in unordered list", style="List Bullet")
    document.add_paragraph("first item in ordered list", style="List Number")
    document.add_picture("python-practice/code.png", width=Inches(1.25))
    records = (
        (3, "101", "Spam"),
        (7, "422", "Eggs"),
        (4, "631", "Spam, spam, eggs, and spam"),
    )
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Qty"
    hdr_cells[1].text = "Id"
    hdr_cells[2].text = "Desc"
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    document.add_page_break()
    document.save("python-practice/demo.docx")


if __name__ == "__main__":
    # create_excel()
    create_word()
